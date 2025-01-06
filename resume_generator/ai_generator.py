from groq import Groq
from dotenv import load_dotenv
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import os

class ResumeGenerator:
    def __init__(self):
        self.document = None
        load_dotenv()
        self.api_key = os.getenv("GROQ_API_KEY")
        if not self.api_key:
            raise Exception("GROQ API key not found")
        self.client = Groq(api_key=self.api_key)

    def generate(self, resume_text, job_desc_text):
        # First, generate optimized content using AI
        optimized_content = self._generate_optimized_content(resume_text, job_desc_text)
        
        # Then, create a properly formatted Word document
        return self._create_document(optimized_content)

    def _generate_optimized_content(self, resume_text, job_desc_text):
        """Generate AI-optimized content for the resume."""
        prompt = f"""
        You are a professional resume writer. Create an optimized version of this resume to better match the job description.
        
        Original Resume:
        {resume_text}
        
        Job Description:
        {job_desc_text}
        
        Instructions:
        1. Keep the same basic information but optimize the wording
        2. Match keywords from the job description
        3. Quantify achievements where possible
        4. Use strong action verbs
        5. Maintain professional formatting
        6. Keep content truthful - don't invent experience
        7. Format in clear sections: Summary, Experience, Skills, Education
        
        Return only the optimized resume text.
        """

        try:
            completion = self.client.chat.completions.create(
                model="mixtral-8x7b-32768",
                messages=[
                    {"role": "system", "content": "You are a professional resume writer."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7,
                max_tokens=2048,
                top_p=1,
                stream=False
            )
            
            return completion.choices[0].message.content.strip()

        except Exception as e:
            print(f"AI generation failed: {str(e)}")
            return resume_text  # Fallback to original text if AI fails

    def _create_document(self, content):
        """Create a properly formatted Word document."""
        self.document = Document()
        
        try:
            # Set up document formatting
            sections = self.document.sections
            for section in sections:
                section.page_width = Inches(8.5)
                section.page_height = Inches(11)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
            
            # Add content with formatting
            title = self.document.add_heading('Optimized Resume', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add content with proper formatting
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    p = self.document.add_paragraph()
                    p.add_run(para.strip())
                    p.space_after = Pt(12)
            
            # Save to BytesIO
            doc_buffer = BytesIO()
            self.document.save(doc_buffer)
            doc_buffer.seek(0)
            
            return doc_buffer.getvalue()
            
        except Exception as e:
            print(f"Document creation failed: {str(e)}")
            return self._generate_basic_document(content)
    
    def _generate_basic_document(self, content):
        """Fallback method to generate a basic document if the main method fails"""
        doc = Document()
        doc.add_paragraph(content)
        
        doc_buffer = BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        return doc_buffer.getvalue()